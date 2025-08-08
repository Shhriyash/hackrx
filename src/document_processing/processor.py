"""
This module provides document processing capabilities using the  library
for parsing PDFs and extracting structured content with page information.
"""

import re
import json
import time
import logging
import tempfile
import os
from itertools import accumulate
from typing import List, Dict, Any, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter

from google import genai
from google.genai import types
import httpx

# For simple text extraction
import fitz  # PyMuPDF - much faster than pymupdf4llm
from docx import Document
import io
from pathlib import Path

from ..interfaces import DocumentProcessor

# Setup logger for this module
logger = logging.getLogger(__name__)


class GeminiFlashProcessor(DocumentProcessor):
    """
    Document processor using Gemini Flash 2.5 for PDF OCR processing.
    
    This processor downloads PDF documents from URLs, processes them with Gemini Flash 2.5
    for OCR (including tables and chart summaries), and returns page-wise content in markdown format.
    """
    
    def __init__(self, api_key: Optional[str] = None, chunk_size: int = 800, chunk_overlap: int = 100):
        """
        Initialize the Gemini Flash processor.
        
        Args:
            api_key: Google AI API key (if None, will use environment variable)
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        logger.info("Initializing GeminiFlashProcessor...")
        start_time = time.time()
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize Gemini client
        client_start = time.time()
        if api_key:
            self.client = genai.Client(api_key=api_key)
        else:
            self.client = genai.Client()  # Will use GOOGLE_API_KEY env var
        client_time = time.time() - client_start
        logger.info(f"Gemini client initialized in {client_time:.3f}s")
            
        # Initialize text splitter
        splitter_start = time.time()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n"],
            is_separator_regex=False
        )
        splitter_time = time.time() - splitter_start
        logger.info(f"Text splitter initialized in {splitter_time:.3f}s")
        
        total_time = time.time() - start_time
        logger.info(f"GeminiFlashProcessor fully initialized in {total_time:.3f}s")
    
    def process_document(self, document_url: str) -> List[Dict[str, Any]]:
        """
        Process a PDF document from URL and return chunks with metadata.
        
        Args:
            document_url: URL to the PDF document
            
        Returns:
            List of dictionaries containing chunk content and metadata
        """
        logger.info(f"Starting Gemini document processing for URL: {document_url}")
        start_time = time.time()
        
        # Download PDF document
        download_start = time.time()
        try:
            logger.info(f"Downloading PDF from: {document_url}")
            response = httpx.get(document_url, timeout=30.0)
            response.raise_for_status()
            doc_data = response.content
            download_time = time.time() - download_start
            logger.info(f"PDF download completed in {download_time:.3f}s - Size: {len(doc_data)} bytes")
        except httpx.HTTPError as e:
            download_time = time.time() - download_start
            logger.error(f"PDF download failed after {download_time:.3f}s: {e}")
            raise ValueError(f"Failed to download PDF from {document_url}: {e}")

        # Simple OCR prompt for faster processing
        ocr_prompt = """
Extract all text content from this PDF document. Return the content in a structured format with page information.
Focus on extracting:
1. All text content verbatim
2. Table content in readable format
3. Brief descriptions of any charts or figures

Return the text in a clean, readable format with clear page breaks indicated.
"""
        
        # Process with Gemini Flash 2.5
        gemini_start = time.time()
        try:
            logger.info("Processing PDF with Gemini Flash 2.5...")
            gemini_response = self.client.models.generate_content(
                model="gemini-2.5-flash-lite",  # Using Flash for faster processing
                contents=[
                    types.Part.from_bytes(
                        data=doc_data,
                        mime_type='application/pdf',
                    ),
                    ocr_prompt
                ]
            )
            
            extracted_text = gemini_response.text
            gemini_time = time.time() - gemini_start
            logger.info(f"Gemini processing completed in {gemini_time:.3f}s - Extracted {len(extracted_text)} characters")
        except Exception as e:
            gemini_time = time.time() - gemini_start
            logger.error(f"Gemini processing failed after {gemini_time:.3f}s: {e}")
            raise ValueError(f"Failed to process PDF with Gemini: {e}")
        
        # Split into chunks
        chunk_start = time.time()
        logger.info("Splitting text into chunks...")
        chunks = self.text_splitter.split_text(extracted_text)
        chunk_time = time.time() - chunk_start
        logger.info(f"Text splitting completed in {chunk_time:.3f}s - Created {len(chunks)} raw chunks")
        
        # Create chunk objects
        process_start = time.time()
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():
                processed_chunks.append({
                    "chunk_id": f"chunk_{i}",
                    "page_number": str(i // 5 + 1),  # Approximate page mapping
                    "content": chunk.strip(),
                    "source_url": document_url,
                    "processor": "gemini_flash_lite"
                })
        
        process_time = time.time() - process_start
        total_time = time.time() - start_time
        
        logger.info(f"Chunk object creation completed in {process_time:.3f}s")
        logger.info(f"Total Gemini document processing completed in {total_time:.3f}s")
        logger.info(f"Timing breakdown - Download: {download_time:.3f}s, Gemini: {gemini_time:.3f}s, Chunking: {chunk_time:.3f}s, Processing: {process_time:.3f}s")
        logger.info(f"Final result: {len(processed_chunks)} processed chunks")
        
        return processed_chunks


class SimpleTextProcessor(DocumentProcessor):
    """
    Simple document processor that extracts text directly from PDFs and DOCX files.
    
    This processor extracts text without OCR, relying on embedded text content.
    Supports both local files and URLs for PDF/DOCX documents.
    """

    def __init__(self, chunk_size: int = 9000, chunk_overlap: int = 300):
        """
        Initialize the Simple Text processor.
        
        Args:
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        logger.info("Initializing SimpleTextProcessor...")
        start_time = time.time()
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        splitter_start = time.time()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " "],
            is_separator_regex=False
        )
        splitter_time = time.time() - splitter_start
        logger.info(f"Text splitter initialized in {splitter_time:.3f}s")
        
        total_time = time.time() - start_time
        logger.info(f"SimpleTextProcessor fully initialized in {total_time:.3f}s")
    
    def process_document(self, document_source: str) -> List[Dict[str, Any]]:
        """
        Process a PDF or DOCX document and return chunks with metadata.
        
        Args:
            document_source: Path to local file or URL to the document
            
        Returns:
            List of dictionaries containing chunk content and metadata
        """
        logger.info(f"Starting simple text extraction for: {document_source}")
        start_time = time.time()
        
        # Determine if source is URL or local file
        is_url = document_source.startswith(('http://', 'https://'))
        
        if is_url:
            # Download document
            download_start = time.time()
            try:
                logger.info(f"Downloading document from: {document_source}")
                response = httpx.get(document_source, timeout=30.0)
                response.raise_for_status()
                doc_data = response.content
                download_time = time.time() - download_start
                logger.info(f"Document download completed in {download_time:.3f}s - Size: {len(doc_data)} bytes")
                
                # Determine file type from content-type or URL
                content_type = response.headers.get('content-type', '').lower()
                if 'pdf' in content_type or document_source.lower().endswith('.pdf'):
                    file_type = 'pdf'
                elif 'word' in content_type or document_source.lower().endswith(('.docx', '.doc')):
                    file_type = 'docx'
                else:
                    # Try to guess from URL extension
                    if document_source.lower().endswith('.pdf'):
                        file_type = 'pdf'
                    elif document_source.lower().endswith(('.docx', '.doc')):
                        file_type = 'docx'
                    else:
                        raise ValueError(f"Unsupported file type. Could not determine if PDF or DOCX from: {document_source}")
                
            except httpx.HTTPError as e:
                download_time = time.time() - download_start
                logger.error(f"Document download failed after {download_time:.3f}s: {e}")
                raise ValueError(f"Failed to download document from {document_source}: {e}")
        else:
            # Local file
            logger.info(f"Processing local file: {document_source}")
            file_path = Path(document_source)
            
            if not file_path.exists():
                raise ValueError(f"File not found: {document_source}")
            
            # Read file
            read_start = time.time()
            with open(file_path, 'rb') as f:
                doc_data = f.read()
            read_time = time.time() - read_start
            logger.info(f"File read completed in {read_time:.3f}s - Size: {len(doc_data)} bytes")
            
            # Determine file type from extension
            if file_path.suffix.lower() == '.pdf':
                file_type = 'pdf'
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                file_type = 'docx'
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}. Only PDF and DOCX are supported.")
            
            download_time = 0  # No download for local files
        
        # Extract text based on file type
        extract_start = time.time()
        if file_type == 'pdf':
            extracted_text = self._extract_pdf_text(doc_data)
        elif file_type == 'docx':
            extracted_text = self._extract_docx_text(doc_data)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        extract_time = time.time() - extract_start
        logger.info(f"Text extraction completed in {extract_time:.3f}s - Extracted {len(extracted_text)} characters")
        
        # Split into chunks
        chunk_start = time.time()
        logger.info("Splitting text into chunks...")
        chunks = self.text_splitter.split_text(extracted_text)
        chunk_time = time.time() - chunk_start
        logger.info(f"Text splitting completed in {chunk_time:.3f}s - Created {len(chunks)} raw chunks")
        
        # Create chunk objects
        process_start = time.time()
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():
                processed_chunks.append({
                    "chunk_id": f"chunk_{i}",
                    "page_number": str(i // 3 + 1),  # Approximate page mapping
                    "content": chunk.strip(),
                    "source": document_source,
                    "file_type": file_type,
                    "processor": "simple_text_extractor"
                })
        
        process_time = time.time() - process_start
        total_time = time.time() - start_time
        
        logger.info(f"Chunk object creation completed in {process_time:.3f}s")
        logger.info(f"Total text extraction completed in {total_time:.3f}s")
        if is_url:
            logger.info(f"Timing breakdown - Download: {download_time:.3f}s, Extract: {extract_time:.3f}s, Chunking: {chunk_time:.3f}s, Processing: {process_time:.3f}s")
        else:
            logger.info(f"Timing breakdown - Read: {read_time:.3f}s, Extract: {extract_time:.3f}s, Chunking: {chunk_time:.3f}s, Processing: {process_time:.3f}s")
        logger.info(f"Final result: {len(processed_chunks)} processed chunks")
        
        return processed_chunks
    
    def _extract_pdf_text(self, pdf_data: bytes) -> str:
        """
        Extract text from PDF data using PyMuPDF.
        
        Args:
            pdf_data: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        logger.info("Extracting text from PDF using PyMuPDF...")
        start_time = time.time()
        
        try:
            # Create a temporary file since fitz needs a file path or BytesIO for stream
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(pdf_data)
                temp_file_path = temp_file.name
            
            try:
                # Open PDF with PyMuPDF
                doc = fitz.open(temp_file_path)
                
                text_parts = []
                page_count = len(doc)
                
                for page_num in range(page_count):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text.strip():
                        # Keep the "--- Page X ---" format for compatibility
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")
                
                doc.close()
                
                extracted_text = "\n\n".join(text_parts)
                extract_time = time.time() - start_time
                logger.info(f"PDF text extraction completed in {extract_time:.3f}s - {page_count} pages processed")
                
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from PDF. The PDF might contain only images or scanned content.")
                
                return extracted_text
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass
            
        except Exception as e:
            extract_time = time.time() - start_time
            logger.error(f"PDF text extraction failed after {extract_time:.3f}s: {e}")
            raise ValueError(f"Failed to extract text from PDF using PyMuPDF: {e}")
    
    def _extract_docx_text(self, docx_data: bytes) -> str:
        """
        Extract text from DOCX data using python-docx.
        
        Args:
            docx_data: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        logger.info("Extracting text from DOCX using python-docx...")
        start_time = time.time()
        
        try:
            # Create document from bytes
            docx_file = io.BytesIO(docx_data)
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            text_parts = []
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    text_parts.append(para_text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("--- Table ---\n" + "\n".join(table_text))
            
            extracted_text = "\n\n".join(text_parts)
            extract_time = time.time() - start_time
            logger.info(f"DOCX text extraction completed in {extract_time:.3f}s - {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables processed")
            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from DOCX document.")
            
            return extracted_text
            
        except Exception as e:
            extract_time = time.time() - start_time
            logger.error(f"DOCX text extraction failed after {extract_time:.3f}s: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
